import streamlit as st
import pandas as pd
import google.generativeai as genai
import json
import time
import os
import zipfile
import re
import hashlib
from docx import Document
from io import BytesIO

# ========================================================
# 1. PAGE CONFIGURATION & SETUP
# ========================================================
st.set_page_config(page_title="ICD-10 Hospital Intelligence (Secure Test)", page_icon="🏥", layout="wide")
st.title("🏥 Hospital Intelligence: case mix analysis")

# ========================================================
# 2. SIDEBAR & USER INSTRUCTIONS
# ========================================================
with st.sidebar:
    # --- INSTRUCTIONS BLOCK ---
    with st.expander("📖 **How to Use**", expanded=True):
        st.markdown("""
        **1. Initialize History**
        Upload your previous history file (if available).
           
        **2. Configure Settings**
        Enter API Keys, Department, and Month.
        
        **3. Upload Current Month Data**
        Upload your raw Excel or CSV file.
           
        **4. Map the Data**
        Select the correct columns (ID, Diagnosis, etc.).
        
        **5. Run Analysis**
        Click **🚀 Secure Run** to generate reports.
        """)
    
    st.divider()
    st.header("⚙️ Configuration")
    
    # --- API KEY INPUT ---
    st.subheader("🔑 Google API Access")
    st.markdown("Use your **Personal/Free** keys for testing.")
    
    c1, c2 = st.columns(2)
    with c1:
        st.link_button("🛠️ Get Key", "https://aistudio.google.com/app/apikey")
    with c2:
        st.link_button("🎥 Tutorial", "https://www.youtube.com/results?search_query=how+to+get+google+gemini+api+key")

    api_keys_input = st.text_area("Enter API Keys (One per line)", height=100, placeholder="Paste keys here...")
    api_keys = [k.strip() for k in api_keys_input.split('\n') if k.strip()]
    
    st.divider()
    
    # --- DEPARTMENT & DATE SELECTION ---
    # Locked list as requested
    dept_options = ["ICU", "CCU", "PICU", "Inpatient Ward", "ER"]
    dept_name = st.selectbox("Department Name", options=dept_options)
    
    col_m, col_y = st.columns(2)
    with col_m: sel_month = st.selectbox("Month", ["Jan", "Feb", "Mar", "Apr", "May", "Jun", "Jul", "Aug", "Sep", "Oct", "Nov", "Dec"])
    with col_y: sel_year = st.selectbox("Year", [str(y) for y in range(2024, 2030)])
    data_month = f"{sel_month} {sel_year}"
    
    # --- API QUOTA TRACKER ---
    st.divider()
    if 'api_calls' not in st.session_state: st.session_state['api_calls'] = 0
    quota_placeholder = st.empty()
    quota_placeholder.metric("API Calls This Session", st.session_state['api_calls'])
    
    # --- RESET BUTTON ---
    st.divider()
    if st.button("🔄 Reset App State", use_container_width=True):
        st.rerun()

# ========================================================
# 3. HELPER FUNCTIONS: SECURITY & AI
# ========================================================

# --- SECURITY: ID HASHING ---
def hash_patient_id(raw_id, salt="TEST_SALT_2026"):
    """
    Converts '12345' -> 'a7x9...' using SHA256. 
    This ensures Real IDs are never sent to the AI.
    """
    raw_str = str(raw_id).strip()
    if not raw_str or raw_str.lower() == 'nan':
        return "UNKNOWN"
    secure_string = raw_str + salt
    return hashlib.sha256(secure_string.encode()).hexdigest()[:12]

# --- TEXT PARSING ---
def extract_json_from_text(text):
    try:
        start = text.find('{')
        end = text.rfind('}') + 1
        if start == -1 or end == 0: return None
        return json.loads(text[start:end])
    except: return None

# --- AI MAPPING ENGINE ---
def get_icd_mapping_optimized(keys_list, unique_diagnoses):
    mapping_dict = {}
    batch_size = 400 
    
    current_key_idx = 0
    models_to_try = ['gemini-2.5-flash', 'gemini-2.5-pro', 'gemini-2.0-flash']
    current_model_idx = 0
    
    genai.configure(api_key=keys_list[current_key_idx])
    model = genai.GenerativeModel(models_to_try[current_model_idx])
    
    total_batches = (len(unique_diagnoses) + batch_size - 1) // batch_size
    pbar = st.progress(0, text="AI Analysis Starting...")
    
    batches = [unique_diagnoses[i:i + batch_size] for i in range(0, len(unique_diagnoses), batch_size)]
    
    for i, batch in enumerate(batches):
        success = False
        max_attempts = len(keys_list) * 2 
        
        for attempt in range(max_attempts):
            try:
                json_payload = json.dumps(batch, default=str)
                prompt = f"""
                Act as a Medical Coder. Map these diagnosis strings to ICD-10.
                Return a JSON object: key = original text, value = {{"Level2": "Block Name", "Level3": "Category Name"}}
                NO markdown. JUST JSON.
                List: {json_payload}
                """
                response = model.generate_content(prompt)
                
                st.session_state['api_calls'] += 1
                quota_placeholder.metric("API Calls This Session", st.session_state['api_calls'])
                
                data = extract_json_from_text(response.text)
                if data: 
                    mapping_dict.update(data)
                    success = True
                    break 
                
            except Exception as e:
                err = str(e)
                if "429" in err or "Quota" in err:
                    # Rotate keys logic
                    if len(keys_list) > 1:
                        current_key_idx = (current_key_idx + 1) % len(keys_list)
                        st.toast(f"Quota hit. Rotating Key...", icon="🔄")
                        genai.configure(api_key=keys_list[current_key_idx])
                        model = genai.GenerativeModel(models_to_try[current_model_idx])
                        time.sleep(1)
                        continue
                    # Downgrade model logic
                    if current_model_idx < len(models_to_try) - 1:
                        current_model_idx += 1
                        st.warning(f"Quota hit. Downgrading to {models_to_try[current_model_idx]}...")
                        model = genai.GenerativeModel(models_to_try[current_model_idx])
                        time.sleep(1)
                        continue
                    st.error("❌ **QUOTA FAILURE** Daily limit reached. Wait 24h or add new key.")
                    st.stop()
                else:
                    time.sleep(1)
        
        if success:
            pbar.progress((i + 1) / total_batches, text=f"Processed batch {i+1}/{total_batches}")
            time.sleep(1)
            
    pbar.empty()
    return mapping_dict

# ========================================================
# 4. HELPER FUNCTIONS: REPORTING & FILES
# ========================================================

def update_master_history(current_stats, dept, month, old_hist=None):
    current_stats['Department'] = dept
    current_stats['Month'] = month
    if old_hist is not None and not old_hist.empty:
        # Remove old entry for this month/dept to avoid duplicates
        old_hist = old_hist[~((old_hist['Department'] == dept) & (old_hist['Month'] == month))]
        return pd.concat([old_hist, current_stats], ignore_index=True)
    return current_stats

# --- RICH FORMATTER (MARKDOWN TO WORD) ---
def write_markdown_to_docx(doc, text):
    """
    Parses Markdown text and applies real Word formatting.
    Handles Headings (#) and Bold (**text**).
    """
    lines = text.split('\n')
    for line in lines:
        line = line.strip()
        if not line:
            continue
            
        # 1. Handle Headings (e.g. ## Title)
        if line.startswith('#'):
            level = min(line.count('#'), 3)
            # Remove hash and bold markers from heading text
            clean_text = line.lstrip('#').strip().replace('**', '').replace('__', '')
            doc.add_heading(clean_text, level=level)
            
        # 2. Handle Lists (e.g. * Item or - Item)
        elif line.startswith(('* ', '- ', '• ')):
            clean_text = line[2:].strip()
            p = doc.add_paragraph(style='List Bullet')
            # Handle Bold inside list items
            parts = re.split(r'(\*\*.*?\*\*)', clean_text)
            for part in parts:
                if part.startswith('**') and part.endswith('**'):
                    run = p.add_run(part[2:-2])
                    run.bold = True
                else:
                    p.add_run(part)
                    
        # 3. Handle Normal Paragraphs
        else:
            p = doc.add_paragraph()
            # Handle Bold inside normal text
            parts = re.split(r'(\*\*.*?\*\*)', line)
            for part in parts:
                if part.startswith('**') and part.endswith('**'):
                    run = p.add_run(part[2:-2])
                    run.bold = True
                else:
                    p.add_run(part)

def generate_report(history_df, month, dept, keys):
    genai.configure(api_key=keys[0])
    model = genai.GenerativeModel('gemini-2.5-flash')
    dept_hist = history_df[history_df['Department'] == dept]
    
    if dept_hist.empty: return None, "No history found for this department."
    
    prompt = f"""
    Act as a **Dual-Role Executive: Strategic Director & Clinical Governance Lead**.
    Your task is to produce a comprehensive **Integrated Case Mix & Strategic Analysis Report** for the '{dept}' department.
    
    **Context:**
    - Current Period: {month}
    - Data Source: Historical ICD-10 Diagnosis Groups (Level 3)
    
    **Dataset (CSV):**
    {dept_hist.to_csv(index=False)}
    
    **Instructions:**
    - Use Markdown formatting (# Heading, **Bold**, * Bullet). NO code blocks.
    
    **Report Structure:**
    1. **Executive Summary** (Business & Clinical)
    2. **Trend & Case Mix Analysis** (Volume & Severity Shifts)
    3. **Operational Efficiency & Clinical Governance** (Skewed LOS & Risk)
    4. **Revenue Integrity** (Profitability vs Cost)
    5. **Integrated Recommendations** (Pathways, Capacity, Coding)
    """
    
    try:
        response = model.generate_content(prompt)
        doc = Document()
        doc.add_heading(f'Strategic Report: {dept} - {month}', 0)
        write_markdown_to_docx(doc, response.text)
        b = BytesIO()
        doc.save(b)
        b.seek(0)
        return b, None 
    except Exception as e: 
        return None, str(e)

# ========================================================
# 5. MAIN APPLICATION LOGIC
# ========================================================

# Initialize Session State
if 'analysis_done' not in st.session_state:
    st.session_state['analysis_done'] = False
if 'zip_buffer' not in st.session_state:
    st.session_state['zip_buffer'] = None
if 'zip_filename' not in st.session_state:
    st.session_state['zip_filename'] = "analysis.zip"

# --- STEP 1: HISTORY INITIALIZATION ---
st.subheader("1. Initialize History (Mandatory)")
st.info("Upload previous history file to maintain trends.")
history_mode = st.radio("Do you have a previous history file?", ["Yes, I have it", "No (Start Fresh)"], horizontal=True)

history_df = pd.DataFrame() 

if history_mode == "Yes, I have it":
    uploaded_history = st.file_uploader("📂 Upload History CSV", type=['csv'], key="hist_up")
    if uploaded_history:
        try:
            history_df = pd.read_csv(uploaded_history)
            st.success(f"✅ History Loaded! Contains {len(history_df)} records.")
        except: st.error("Invalid CSV.")
    else:
        st.warning("⚠️ Upload history to proceed.")
        st.stop()
else:
    st.success("✅ Starting with a fresh history database.")

st.divider()

# --- STEP 2: DATA UPLOAD & MAPPING ---
st.subheader("2. Upload Current Month Data")
uploaded_file = st.file_uploader("Upload Raw Hospital Data (CSV/Excel)", type=['csv', 'xlsx'], key="data_up")

if uploaded_file and api_keys:
    try:
        df = pd.read_csv(uploaded_file) if uploaded_file.name.endswith('.csv') else pd.read_excel(uploaded_file)
        st.dataframe(df.head(3))
        
        st.divider()
        st.subheader("3. Map Columns & Secure Run")
        cols = list(df.columns)
        c1, c2, c3 = st.columns(3)
        col_id = c1.selectbox("Patient ID", cols, index=0)
        col_diag = c2.selectbox("Diagnosis", cols, index=min(1, len(cols)-1))
        col_rev = c3.selectbox("Revenue", cols, index=min(2, len(cols)-1))
        col_adm = c1.selectbox("Admission Date", cols, index=min(3, len(cols)-1))
        col_disch = c2.selectbox("Discharge Date", cols, index=min(4, len(cols)-1))
        
        if st.button("🚀 Secure Run (Gemini)", type="primary"):
            with st.status("Processing Securely...", expanded=True):
                # 1. CLEAN & ANONYMIZE
                st.write("🔒 Anonymizing Patient IDs (SHA-256)...")
                wdf = df.rename(columns={col_id:'ID', col_diag:'DIAG', col_rev:'REV', col_adm:'ADM', col_disch:'DISCH'})
                
                # --- APPLY ENCRYPTION ---
                wdf['Original_ID'] = wdf['ID'] 
                wdf['ID'] = wdf['ID'].apply(hash_patient_id) 
                
                wdf = wdf.dropna(subset=['DIAG']) 
                wdf['DIAG'] = wdf['DIAG'].astype(str) 
                wdf = wdf[wdf['DIAG'].str.strip() != ''] 
                wdf['ADM'] = pd.to_datetime(wdf['ADM'], errors='coerce')
                wdf['DISCH'] = pd.to_datetime(wdf['DISCH'], errors='coerce')
                wdf['LOS'] = (wdf['DISCH'] - wdf['ADM']).dt.days.fillna(0).clip(lower=0)
                wdf['REV'] = pd.to_numeric(wdf['REV'], errors='coerce').fillna(0)
                
                # 2. MAP
                st.write(f"🧠 AI Mapping ({len(wdf)} records)...")
                unique = wdf['DIAG'].unique().tolist()
                mapping = get_icd_mapping_optimized(api_keys, unique)
                
                if not mapping:
                    st.error("AI Mapping Failed.")
                    st.stop()

                wdf['Level3'] = wdf['DIAG'].map(lambda x: mapping.get(x, {}).get("Level3", "Unmapped"))
                
                # 3. STATS
                stats = wdf.groupby('Level3').agg({'DIAG':'count', 'LOS':['mean','median'], 'REV':['mean','median']}).reset_index()
                stats.columns = ['Diagnosis Group', 'Cases', 'Mean LOS', 'Median LOS', 'Mean Rev', 'Median Rev']
                stats['Skew Status'] = stats.apply(lambda x: 'Skewed' if x['Mean LOS'] > (x['Median LOS']*1.5) else 'Normal', axis=1)
                
                # 4. HISTORY
                st.write("📜 Updating History...")
                full_hist = update_master_history(stats, dept_name, data_month, history_df)
                
                # 5. REPORT
                st.write("📝 Generating Strategic Report...")
                report_doc, report_error = generate_report(full_hist, data_month, dept_name, api_keys)
                
                # 6. PACKAGING & NAMING
                safe_dept = dept_name.replace(" ", "_").replace("/", "-")
                safe_month = data_month.replace(" ", "_")
                
                history_filename = f"{safe_month}-{safe_dept}-history.csv"
                zip_filename = f"Hospital_Analysis_{safe_month}-{safe_dept}.zip"
                
                audit_buffer = BytesIO()
                wdf.to_excel(audit_buffer, index=False)
                
                zip_buffer = BytesIO()
                with zipfile.ZipFile(zip_buffer, "w") as zf:
                    # History (Mandatory)
                    zf.writestr(history_filename, full_hist.to_csv(index=False))
                    # Report (Docx)
                    if report_doc:
                        zf.writestr(f"Report_{safe_month}-{safe_dept}.docx", report_doc.getvalue())
                    # Audit (Excel)
                    zf.writestr(f"Audit_Encrypted_{safe_month}-{safe_dept}.xlsx", audit_buffer.getvalue())
                    
                zip_buffer.seek(0)
                
                st.session_state['zip_buffer'] = zip_buffer
                st.session_state['analysis_done'] = True
                st.session_state['full_hist'] = full_hist
                st.session_state['zip_filename'] = zip_filename
                
            st.success("Analysis Complete!")
            st.rerun()

    except Exception as e:
        st.error(f"Critical App Error: {e}")

# ========================================================
# 6. RESULTS & DOWNLOAD DISPLAY
# ========================================================
if st.session_state.get('analysis_done') and st.session_state.get('zip_buffer'):
    st.divider()
    st.subheader("✅ Results Ready")
    st.warning("⬇️ **STEP 4: Download Files**")
    
    st.download_button(
        label="📦 DOWNLOAD ALL FILES (ZIP)",
        data=st.session_state['zip_buffer'],
        file_name=st.session_state['zip_filename'],
        mime="application/zip",
        type="primary",
        use_container_width=True
    )
    
    st.divider()
    col_reset, _ = st.columns([1,3])
    with col_reset:
        if st.button("🔄 Start New Analysis (Reset Page)", type="secondary"):
            for key in list(st.session_state.keys()):
                del st.session_state[key]
            st.rerun()

    st.divider()
    if 'full_hist' in st.session_state:
        st.write("📊 **History Preview:**")
        st.dataframe(st.session_state['full_hist'].tail(10))
